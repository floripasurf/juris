"""Markdown → .docx export (entrega): lawyers finish drafts in Word.

Covers the converter (valid docx, headings/paragraphs/bold/lists) and the
tenant-gated endpoint that serves it from disk (full fidelity, no truncation)
or from posted text (the edited filing draft).
"""

from __future__ import annotations

import io
import json

import pytest
from fastapi.testclient import TestClient

from juris.web.app import app
from juris.web.auth import default_registry, hash_api_key

MINUTA = """# EXCELENTÍSSIMO SENHOR DOUTOR JUIZ

**AUTOR**, já qualificado, vem apresentar CONTESTAÇÃO.

## DOS FATOS

O réu nunca foi notificado.

- Primeiro ponto
- Segundo ponto

## DO PEDIDO

Requer a improcedência.
"""


def test_markdown_to_docx_produces_readable_docx() -> None:
    from docx import Document

    from juris.web.export import markdown_to_docx

    data = markdown_to_docx(MINUTA)
    assert data[:2] == b"PK"  # docx is a zip
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "EXCELENTÍSSIMO" in text
    assert "DOS FATOS" in text
    assert "Primeiro ponto" in text
    # bold run survived
    bold_runs = [r.text for p in doc.paragraphs for r in p.runs if r.bold]
    assert any("AUTOR" in t for t in bold_runs)


@pytest.fixture
def tenant_env(monkeypatch, tmp_path):
    tenants = tmp_path / "tenants.json"
    tenants.write_text(json.dumps({"escritorio-a": hash_api_key("key-a")}), encoding="utf-8")
    monkeypatch.setenv("JURIS_REQUIRE_TENANTS", "1")
    monkeypatch.setenv("JURIS_TENANTS_FILE", str(tenants))
    monkeypatch.setenv("JURIS_HOME", str(tmp_path))
    monkeypatch.setenv("JURIS_OUT_ROOT", str(tmp_path / "out"))
    default_registry.cache_clear()
    yield {"headers": {"X-API-Key": "key-a"}, "out_root": tmp_path / "out"}
    default_registry.cache_clear()


class TestExportEndpoint:
    def test_export_from_posted_markdown(self, tenant_env) -> None:
        client = TestClient(app)
        r = client.post(
            "/api/export/docx",
            json={"markdown": MINUTA, "filename": "contestacao.docx"},
            headers=tenant_env["headers"],
        )
        assert r.status_code == 200
        assert r.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert "contestacao.docx" in r.headers.get("content-disposition", "")
        assert r.content[:2] == b"PK"

    def test_export_requires_tenant(self, tenant_env) -> None:
        client = TestClient(app)
        assert client.post("/api/export/docx", json={"markdown": "x"}).status_code == 401

    def test_export_rejects_empty(self, tenant_env) -> None:
        client = TestClient(app)
        r = client.post("/api/export/docx", json={"markdown": "   "}, headers=tenant_env["headers"])
        assert r.status_code == 400

    def test_export_from_disk_is_tenant_confined(self, tenant_env, tmp_path) -> None:
        # path traversal / escaping the tenant dir must be refused
        client = TestClient(app)
        r = client.post(
            "/api/export/docx",
            json={"output_dir": "../../etc", "name": "passwd"},
            headers=tenant_env["headers"],
        )
        assert r.status_code in (400, 404)


class TestDeliveryUI:
    """Static SPA pins for the entrega batch (download + rendered preview)."""

    def test_index_ships_docx_download_and_markdown_preview(self) -> None:
        from pathlib import Path

        html = (
            Path(__file__).resolve().parents[3] / "src" / "juris" / "web" / "static" / "index.html"
        ).read_text(encoding="utf-8")
        # download helper + wired buttons
        assert "function downloadDocx(" in html
        assert 'id="artifact-download"' in html
        assert 'id="fl_download_docx"' in html
        assert "/api/export/docx" in html
        # safe rendered preview + toggle
        assert "function renderMarkdownSafe(" in html
        assert 'data-mode="formatado"' in html
        assert 'data-mode="markdown"' in html

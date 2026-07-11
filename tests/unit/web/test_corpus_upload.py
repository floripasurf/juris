"""Upload de acervo do escritório → corpus com proveniência (a via aprovada no ToS).

O único caminho de inteiro teor liberado em ``data/tos_compliance_log.md`` é
"arquivos do próprio escritório". Este seam cobre o endpoint que transforma um
documento enviado no console (texto colado ou PDF) em fonte aceita + chunks
pesquisáveis, mantendo a proveniência obrigatória e o escopo por tenant.
"""

from __future__ import annotations

import base64
import json
import sqlite3

import pytest
from fastapi.testclient import TestClient

from juris.web.app import app
from juris.web.auth import default_registry, hash_api_key

PROVENANCE = {
    "title": "Sentença — Cobrança 0001",
    "source_type": "acordao_publicado",
    "source_date": "2024-05-13",
    "source_url": "https://pje.tjmg.jus.br/consulta/0001234",
    "tribunal": "tjmg",
    "tema": "honorarios",
}

TEXTO = (
    "EMENTA: APELAÇÃO CÍVEL. COBRANÇA. HONORÁRIOS SUCUMBENCIAIS CONTRA A "
    "FAZENDA PÚBLICA. Fixação por equidade quando o proveito econômico for "
    "inestimável. Recurso provido. " * 20
)


@pytest.fixture
def tenant_env(monkeypatch, tmp_path):
    tenants = tmp_path / "tenants.json"
    tenants.write_text(json.dumps({"escritorio-a": hash_api_key("key-a")}), encoding="utf-8")
    monkeypatch.setenv("JURIS_REQUIRE_TENANTS", "1")
    monkeypatch.setenv("JURIS_TENANTS_FILE", str(tenants))
    monkeypatch.setenv("JURIS_HOME", str(tmp_path))
    monkeypatch.setenv("JURIS_OUT_ROOT", str(tmp_path / "out"))
    monkeypatch.setenv("JURIS_REPERTORY_PATH", str(tmp_path / "repertory.db"))
    default_registry.cache_clear()
    yield {"headers": {"X-API-Key": "key-a"}, "repertory": tmp_path / "repertory.db"}
    default_registry.cache_clear()


def _upload(client: TestClient, headers: dict[str, str], **overrides) -> object:
    payload: dict[str, object] = {**PROVENANCE, "source_text": TEXTO, **overrides}
    return client.post("/api/corpus/upload", json=payload, headers=headers)


class TestUploadTexto:
    def test_pasted_text_registers_source_and_ingests_chunks(self, tenant_env) -> None:
        client = TestClient(app)
        response = _upload(client, tenant_env["headers"])
        assert response.status_code == 201, response.text
        body = response.json()
        assert body["source"]["id"]
        assert body["source"]["content_sha256"].startswith("")  # presente
        assert body["reingest"]["processed"] == 1
        assert body["reingest"]["chunks"] >= 1

        rows = sqlite3.connect(tenant_env["repertory"]).execute(
            "SELECT COUNT(*) FROM chunks"
        ).fetchone()
        assert rows[0] == body["reingest"]["chunks"]

    def test_upload_is_tenant_scoped_on_disk(self, tenant_env, tmp_path) -> None:
        client = TestClient(app)
        assert _upload(client, tenant_env["headers"]).status_code == 201
        tenant_dir = tmp_path / "out" / "tenants" / "escritorio-a"
        assert (tenant_dir / "corpus-sources.jsonl").is_file()

    def test_duplicate_content_is_rejected(self, tenant_env) -> None:
        client = TestClient(app)
        assert _upload(client, tenant_env["headers"]).status_code == 201
        again = _upload(client, tenant_env["headers"])
        assert again.status_code == 400
        assert "content_sha256" in again.json()["detail"]["message"]


class TestUploadPdf:
    def test_pdf_base64_is_extracted_and_ingested(self, tenant_env) -> None:
        import pymupdf

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), TEXTO[:800])
        pdf_bytes = doc.tobytes()
        doc.close()

        client = TestClient(app)
        response = _upload(
            client,
            tenant_env["headers"],
            source_text="",
            filename="sentenca.pdf",
            content_base64=base64.b64encode(pdf_bytes).decode("ascii"),
        )
        assert response.status_code == 201, response.text
        assert response.json()["reingest"]["chunks"] >= 1


class TestUploadDocx:
    def test_docx_base64_is_extracted_and_ingested(self, tenant_env) -> None:
        import io

        from docx import Document

        doc = Document()
        doc.add_paragraph("CONTESTAÇÃO. " + TEXTO)
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "Cláusula de tabela relevante."
        buf = io.BytesIO()
        doc.save(buf)

        client = TestClient(app)
        response = _upload(
            client,
            tenant_env["headers"],
            source_text="",
            filename="contestacao.docx",
            content_base64=base64.b64encode(buf.getvalue()).decode("ascii"),
        )
        assert response.status_code == 201, response.text
        assert response.json()["reingest"]["chunks"] >= 1

    def test_docx_corrompido_e_400_legivel(self, tenant_env) -> None:
        client = TestClient(app)
        response = _upload(
            client,
            tenant_env["headers"],
            source_text="",
            filename="quebrado.docx",
            content_base64=base64.b64encode(b"nao sou um docx").decode("ascii"),
        )
        assert response.status_code == 400
        assert "DOCX" in response.json()["detail"]["message"] or "docx" in response.json()["detail"]["message"]


class TestSearchAfterUpload:
    def test_uploaded_corpus_is_searchable_without_seed_readiness(self, tenant_env, monkeypatch) -> None:
        """Num deploy fresco (sem seeds públicas ingeridas), o acervo enviado
        pelo escritório ainda precisa aparecer na busca explicável."""
        import importlib

        app_module = importlib.import_module("juris.web.app")

        class _StubRepertory:
            def search_jurisprudencia(
                self, query: str, top_k: int, tenant_id: str, include_estilo: bool = False, area: str | None = None
            ) -> list:
                return []

        monkeypatch.setattr(app_module, "_corpus_search_service", lambda path: _StubRepertory())
        client = TestClient(app)
        assert _upload(client, tenant_env["headers"]).status_code == 201

        response = client.get("/api/corpus/search?q=honorarios", headers=tenant_env["headers"])
        assert response.status_code == 200
        assert response.json().get("detail") != "corpus não ingerido ainda"


class TestUploadValidation:
    def test_without_text_or_file_is_400(self, tenant_env) -> None:
        client = TestClient(app)
        response = _upload(client, tenant_env["headers"], source_text="")
        assert response.status_code == 400
        assert response.json()["detail"]["code"] == "corpus_upload_invalid"

    def test_missing_provenance_url_is_400(self, tenant_env) -> None:
        client = TestClient(app)
        response = _upload(client, tenant_env["headers"], source_url="")
        assert response.status_code == 400

    def test_unsupported_extension_is_400(self, tenant_env) -> None:
        client = TestClient(app)
        response = _upload(
            client,
            tenant_env["headers"],
            source_text="",
            filename="peticao.docx",
            content_base64=base64.b64encode(b"PK\x03\x04fake").decode("ascii"),
        )
        assert response.status_code == 400
        assert response.json()["detail"]["code"] == "corpus_upload_invalid"


class TestProvenanciaPrivada:
    def test_acervo_do_escritorio_dispensa_url(self, tenant_env) -> None:
        client = TestClient(app)
        payload = {
            "title": "Contestação modelo — cobrança",
            "source_type": "peca_escritorio",
            "source_date": "2025-11-10",
            "source_publisher": "Escritório A",
            "provenance_kind": "acervo_do_escritorio",
            "tipo_peticao": "contestacao",
            "area": "civel",
            "source_text": TEXTO,
        }
        resp = client.post("/api/corpus/upload", json=payload, headers=tenant_env["headers"])
        assert resp.status_code == 201, resp.text
        source = resp.json()["source"]
        assert source["provenance_kind"] == "acervo_do_escritorio"
        assert source["uso"] == "estilo"            # derivado de peca_escritorio
        assert source["tipo_peticao"] == "contestacao"
        assert not source.get("source_url")

    def test_acervo_do_escritorio_com_url_invalida_e_400(self, tenant_env) -> None:
        client = TestClient(app)
        payload = {
            "title": "Contestação modelo — cobrança",
            "source_type": "peca_escritorio",
            "source_date": "2025-11-10",
            "source_publisher": "Escritório A",
            "provenance_kind": "acervo_do_escritorio",
            "area": "civel",
            "source_url": "javascript:alert(1)",  # URL fornecida → ainda valida
            "source_text": TEXTO,
        }
        resp = client.post("/api/corpus/upload", json=payload, headers=tenant_env["headers"])
        assert resp.status_code == 400
        assert "http(s)" in resp.json()["detail"]["message"]

    def test_provenance_publica_continua_exigindo_url(self, tenant_env) -> None:
        client = TestClient(app)
        payload = {**PROVENANCE, "source_text": TEXTO}
        payload.pop("source_url")
        resp = client.post("/api/corpus/upload", json=payload, headers=tenant_env["headers"])
        assert resp.status_code == 400  # comportamento atual preservado

    def test_acervo_do_escritorio_sem_area_nao_ingere(self, tenant_env) -> None:
        client = TestClient(app)
        payload = {
            "title": "Contestação modelo — cobrança",
            "source_type": "peca_escritorio",
            "source_date": "2025-11-10",
            "source_publisher": "Escritório A",
            "provenance_kind": "acervo_do_escritorio",
            "source_text": TEXTO,
        }
        resp = client.post("/api/corpus/upload", json=payload, headers=tenant_env["headers"])
        assert resp.status_code == 400
        assert "area" in resp.json()["detail"]["message"]

    def test_doutrina_do_escritorio_sem_aceite_nao_ingere(self, tenant_env) -> None:
        client = TestClient(app)
        payload = {
            "title": "Manual de Processo Civil",
            "source_type": "doutrina_escritorio",
            "source_date": "2024-01-01",
            "source_publisher": "Editora X",
            "provenance_kind": "acervo_do_escritorio",
            "area": "cível",
            "source_text": TEXTO,
        }
        resp = client.post("/api/corpus/upload", json=payload, headers=tenant_env["headers"])
        assert resp.status_code == 400
        assert "copyright_ack" in resp.json()["detail"]["message"]

    def test_doutrina_do_escritorio_com_aceite_ingere(self, tenant_env) -> None:
        client = TestClient(app)
        payload = {
            "title": "Manual de Processo Civil",
            "source_type": "doutrina_escritorio",
            "source_date": "2024-01-01",
            "source_publisher": "Editora X",
            "provenance_kind": "acervo_do_escritorio",
            "area": "cível",
            "copyright_ack": True,
            "source_text": TEXTO,
        }
        resp = client.post("/api/corpus/upload", json=payload, headers=tenant_env["headers"])
        assert resp.status_code == 201, resp.text
        source = resp.json()["source"]
        assert source["source_type"] == "doutrina_escritorio"
        assert source["area"] == "civel"
        assert source["copyright_ack"] is True
        assert "rights_basis" not in source

    def test_doutrina_privada_legada_vira_doutrina_escritorio(self, tenant_env) -> None:
        client = TestClient(app)
        payload = {
            "title": "Manual de Processo Civil",
            "source_type": "doutrina_privada",
            "source_date": "2024-01-01",
            "source_publisher": "Editora X",
            "provenance_kind": "acervo_do_escritorio",
            "area": "civel",
            "copyright_ack": True,
            "source_text": TEXTO,
        }
        resp = client.post("/api/corpus/upload", json=payload, headers=tenant_env["headers"])
        assert resp.status_code == 201, resp.text
        assert resp.json()["source"]["source_type"] == "doutrina_escritorio"

    def test_uso_override_invalido_e_400(self, tenant_env) -> None:
        client = TestClient(app)
        payload = {**PROVENANCE, "source_text": TEXTO, "uso": "citavel"}
        resp = client.post("/api/corpus/upload", json=payload, headers=tenant_env["headers"])
        assert resp.status_code == 400

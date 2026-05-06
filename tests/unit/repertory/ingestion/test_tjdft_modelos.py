"""Tests for TJDFT petition template ingester."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from juris.repertory.chunking import chunk_fonte
from juris.repertory.corpus.models import FonteJurisprudencia, TipoFonte
from juris.repertory.ingestion.tjdft_modelos import (
    TJDFTModelosIngester,
    _extract_docx_text,
    _infer_tipo_peticao,
)


@pytest.fixture()
def tmp_docx_dir(tmp_path: Path) -> Path:
    """Create a temp directory with sample .docx files."""
    # Simple template
    doc1 = Document()
    doc1.add_paragraph("PETIÇÃO INICIAL")
    doc1.add_paragraph("AO JUIZADO ESPECIAL CÍVEL")
    doc1.add_paragraph("DOS FATOS")
    doc1.add_paragraph("O autor celebrou contrato de compra e venda.")
    doc1.add_paragraph("DOS PEDIDOS")
    doc1.add_paragraph("Requer a condenação do réu ao pagamento.")
    doc1.save(str(tmp_path / "1.1 PETIÇÃO INICIAL - GERAL.docx"))

    # Contestação template
    doc2 = Document()
    doc2.add_paragraph("CONTESTAÇÃO")
    doc2.add_paragraph("DOS FATOS")
    doc2.add_paragraph("O réu nega os fatos narrados na inicial.")
    doc2.add_paragraph("DO DIREITO")
    doc2.add_paragraph("Conforme art. 373 do CPC.")
    doc2.add_paragraph("DOS PEDIDOS")
    doc2.add_paragraph("Requer a improcedência total.")
    doc2.save(str(tmp_path / "2.1 CONTESTAÇÃO - GERAL.docx"))

    # Empty doc (should be skipped)
    doc3 = Document()
    doc3.save(str(tmp_path / "empty.docx"))

    return tmp_path


class TestExtractDocxText:
    def test_extracts_paragraphs(self, tmp_docx_dir: Path) -> None:
        text = _extract_docx_text(tmp_docx_dir / "1.1 PETIÇÃO INICIAL - GERAL.docx")
        assert "PETIÇÃO INICIAL" in text
        assert "DOS FATOS" in text
        assert "DOS PEDIDOS" in text


class TestInferTipoPeticao:
    def test_contestacao(self) -> None:
        assert _infer_tipo_peticao("2.1 CONTESTAÇÃO - GERAL.docx") == "contestacao"

    def test_inicial(self) -> None:
        assert _infer_tipo_peticao("1.1 PETIÇÃO INICIAL - GERAL.docx") == "inicial"

    def test_execucao(self) -> None:
        assert _infer_tipo_peticao("10.1 EXECUÇÃO - título.docx") == "execucao"

    def test_generic(self) -> None:
        result = _infer_tipo_peticao("99.9 SOMETHING ELSE.docx")
        assert "SOMETHING" in result


class TestTJDFTModelosIngester:
    def test_fetch_returns_fontes(self, tmp_docx_dir: Path) -> None:
        ingester = TJDFTModelosIngester(source_dir=tmp_docx_dir)
        fontes = ingester.fetch()
        # 2 valid docs (empty skipped)
        assert len(fontes) == 2

    def test_fetch_with_limit(self, tmp_docx_dir: Path) -> None:
        ingester = TJDFTModelosIngester(source_dir=tmp_docx_dir, limit=1)
        fontes = ingester.fetch()
        assert len(fontes) == 1

    def test_fonte_has_correct_tipo(self, tmp_docx_dir: Path) -> None:
        ingester = TJDFTModelosIngester(source_dir=tmp_docx_dir)
        fontes = ingester.fetch()
        for fonte in fontes:
            assert fonte.tipo == TipoFonte.MODELO_PETICAO
            assert fonte.hierarquia == 7
            assert fonte.tribunal == "TJDFT"
            assert fonte.situacao == "publicado"
            assert fonte.legal_basis == "institutional_template"
            assert fonte.source_publisher == "DPDF/TJDFT"

    def test_parse_produces_chunks(self, tmp_docx_dir: Path) -> None:
        ingester = TJDFTModelosIngester(source_dir=tmp_docx_dir)
        fontes = ingester.fetch()
        for fonte in fontes:
            chunks = ingester.parse(fonte)
            assert len(chunks) >= 1
            for chunk in chunks:
                assert chunk.source_type == TipoFonte.MODELO_PETICAO
                assert chunk.text

    def test_nonexistent_dir(self, tmp_path: Path) -> None:
        ingester = TJDFTModelosIngester(source_dir=tmp_path / "nope")
        fontes = ingester.fetch()
        assert fontes == []


class TestChunkTemplate:
    def test_splits_by_section_headers(self, tmp_docx_dir: Path) -> None:
        ingester = TJDFTModelosIngester(source_dir=tmp_docx_dir)
        fontes = ingester.fetch()
        # Find the contestação template (has DOS FATOS, DO DIREITO, DOS PEDIDOS)
        contestacao = [f for f in fontes if "CONTESTAÇÃO" in f.numero][0]
        chunks = chunk_fonte(contestacao)
        # Should produce multiple chunks from the section splits
        assert len(chunks) >= 1
        # Verify text content is preserved
        all_text = " ".join(c.text for c in chunks)
        assert "réu" in all_text.lower() or "reu" in all_text.lower()

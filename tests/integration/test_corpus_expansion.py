"""Integration test for Sprint 13 corpus expansion.

Verifies end-to-end: type system → ingester → chunking → query
for all new TipoFonte entries.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from juris.repertory.chunking import chunk_fonte
from juris.repertory.corpus.models import TIPO_HIERARQUIA, FonteJurisprudencia, TipoFonte
from juris.repertory.corpus.status import ACTIVE_SITUACOES, is_active


class TestNewTipoFonteIntegration:
    """Verify new TipoFonte members work end-to-end."""

    def test_all_new_tipos_in_hierarchy_map(self) -> None:
        new_tipos = [
            TipoFonte.MODELO_PETICAO,
            TipoFonte.DOUTRINA_PD,
            TipoFonte.NOTICIA_TRIBUNAL,
            TipoFonte.ACORDAO_LANDMARK,
            TipoFonte.ACORDAO_PUBLICADO,
        ]
        for tipo in new_tipos:
            assert tipo in TIPO_HIERARQUIA, f"{tipo} missing from TIPO_HIERARQUIA"

    def test_all_new_tipos_in_active_situacoes(self) -> None:
        new_tipos = [
            TipoFonte.MODELO_PETICAO,
            TipoFonte.DOUTRINA_PD,
            TipoFonte.NOTICIA_TRIBUNAL,
            TipoFonte.ACORDAO_LANDMARK,
            TipoFonte.ACORDAO_PUBLICADO,
        ]
        for tipo in new_tipos:
            assert tipo in ACTIVE_SITUACOES, f"{tipo} missing from ACTIVE_SITUACOES"

    def test_publicado_is_active_for_new_tipos(self) -> None:
        assert is_active(TipoFonte.MODELO_PETICAO, "publicado")
        assert is_active(TipoFonte.DOUTRINA_PD, "publicado")
        assert is_active(TipoFonte.NOTICIA_TRIBUNAL, "publicado")
        assert is_active(TipoFonte.ACORDAO_LANDMARK, "publicado")
        assert is_active(TipoFonte.ACORDAO_PUBLICADO, "publicado")


class TestModeloPeticaoEndToEnd:
    """MODELO_PETICAO: ingest → chunk → verify retrieval shape."""

    def test_modelo_peticao_chunks_correctly(self) -> None:
        fonte = FonteJurisprudencia(
            id="modelo_peticao_TJDFT_test",
            tribunal="TJDFT",
            tipo=TipoFonte.MODELO_PETICAO,
            numero="test.docx",
            ementa="PETIÇÃO INICIAL",
            texto_integral=(
                "PETIÇÃO INICIAL\n\n"
                "AO JUIZADO ESPECIAL CÍVEL\n\n"
                "DOS FATOS\n\n"
                "O autor celebrou contrato.\n\n"
                "DOS PEDIDOS\n\n"
                "Requer a condenação do réu."
            ),
            situacao="publicado",
            hierarquia=7,
            legal_basis="institutional_template",
            source_publisher="DPDF/TJDFT",
        )
        chunks = chunk_fonte(fonte)
        assert len(chunks) >= 1
        assert all(c.source_type == TipoFonte.MODELO_PETICAO for c in chunks)

    def test_modelo_peticao_provenance_fields(self) -> None:
        fonte = FonteJurisprudencia(
            id="test",
            tribunal="TJDFT",
            tipo=TipoFonte.MODELO_PETICAO,
            numero="1",
            ementa="test",
            hierarquia=7,
            source_url="https://tjdft.jus.br",
            source_publisher="DPDF/TJDFT",
            legal_basis="institutional_template",
        )
        assert fonte.source_url == "https://tjdft.jus.br"
        assert fonte.source_publisher == "DPDF/TJDFT"
        assert fonte.legal_basis == "institutional_template"


class TestAcordaoLandmarkEndToEnd:
    """ACORDAO_LANDMARK: ingest → chunk → verify correct strategy."""

    def test_landmark_uses_acordao_chunking(self) -> None:
        fonte = FonteJurisprudencia(
            id="acordao_landmark_STF_ADC_12",
            tribunal="STF",
            tipo=TipoFonte.ACORDAO_LANDMARK,
            numero="ADC_12",
            ementa="EMENTA: Teste de acórdão landmark.",
            texto_integral=(
                "EMENTA\nTeste de ementa.\n\n"
                "RELATÓRIO\nO relator...\n\n"
                "VOTO\nVoto do ministro..."
            ),
            hierarquia=3,
            situacao="publicado",
            legal_basis="government_publication",
        )
        chunks = chunk_fonte(fonte)
        assert len(chunks) >= 2  # ementa + sections
        assert all(c.source_type == TipoFonte.ACORDAO_LANDMARK for c in chunks)


class TestNoticiaTribunalEndToEnd:
    """NOTICIA_TRIBUNAL: chunk as single article."""

    def test_noticia_single_chunk(self) -> None:
        fonte = FonteJurisprudencia(
            id="noticia_tribunal_STF_test",
            tribunal="STF",
            tipo=TipoFonte.NOTICIA_TRIBUNAL,
            numero="info_1",
            ementa="STF decide sobre constitucionalidade.",
            hierarquia=7,
            situacao="publicado",
        )
        chunks = chunk_fonte(fonte)
        assert len(chunks) == 1


class TestDoutrinaPDEndToEnd:
    """DOUTRINA_PD: paragraph-based chunking."""

    def test_doutrina_paragraph_chunking(self) -> None:
        paragraphs = "\n\n".join([f"Parágrafo {i} do texto doutrinário." for i in range(5)])
        fonte = FonteJurisprudencia(
            id="doutrina_pd_test",
            tribunal="STF",
            tipo=TipoFonte.DOUTRINA_PD,
            numero="test",
            ementa="Doutrina teste",
            texto_integral=paragraphs,
            hierarquia=6,
            situacao="publicado",
            legal_basis="government_publication",
        )
        chunks = chunk_fonte(fonte)
        assert len(chunks) >= 1
        assert all(c.source_type == TipoFonte.DOUTRINA_PD for c in chunks)


class TestTJDFTIngesterIntegration:
    """Full pipeline test with temp .docx files."""

    def test_tjdft_pipeline(self, tmp_path: Path) -> None:
        from juris.repertory.ingestion.tjdft_modelos import TJDFTModelosIngester

        doc = Document()
        doc.add_paragraph("CONTESTAÇÃO")
        doc.add_paragraph("DOS FATOS")
        doc.add_paragraph("O réu contesta os pedidos da inicial.")
        doc.add_paragraph("DOS PEDIDOS")
        doc.add_paragraph("Requer a improcedência.")
        doc.save(str(tmp_path / "contestacao.docx"))

        ingester = TJDFTModelosIngester(source_dir=tmp_path)
        fontes = ingester.fetch()
        assert len(fontes) == 1

        fonte = fontes[0]
        assert fonte.tipo == TipoFonte.MODELO_PETICAO
        assert fonte.hierarquia == 7

        chunks = chunk_fonte(fonte)
        assert len(chunks) >= 1
        all_text = " ".join(c.text for c in chunks)
        assert "réu" in all_text.lower() or "reu" in all_text.lower()

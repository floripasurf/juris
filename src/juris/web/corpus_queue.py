"""Pilot-driven corpus queue.

Turns real pilot gaps into an auditable queue of accepted corpus sources. This
module deliberately records provenance and coverage metadata before any vector
reingestion happens, so the corpus grows from lawyer-approved evidence rather
than ad hoc files.
"""

from __future__ import annotations

import base64
import binascii
import hashlib
import json
import re
import uuid
from dataclasses import dataclass, replace
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Any, cast
from urllib.parse import urlparse, urlunparse

from juris.core.paths import ensure_private_dir, restrict_file
from juris.repertory.corpus.models import RIGHTS_BASIS_VALUES, TipoFonte, resolve_uso
from juris.web.pilot_feedback import list_feedback

_FILENAME = "corpus-sources.jsonl"
_TEXT_DIR = "corpus-source-text"
_SHA256_RE = re.compile(r"^[0-9a-f]{64}$")
_REINGEST_SOURCE_ERROR = "Falha ao reingerir esta fonte. Verifique texto, tipo e corpus local."


@dataclass(frozen=True, slots=True)
class ReingestReport:
    processed: int
    chunks: int
    errors: list[dict[str, str]]

    def to_dict(self) -> dict[str, object]:
        return {
            "processed": self.processed,
            "chunks": self.chunks,
            "errors": self.errors,
        }


def sources_path(root: Path) -> Path:
    return root / _FILENAME


def corpus_candidates(root: Path) -> list[dict[str, object]]:
    """Feedback records marked as useful for corpus expansion."""
    accepted_cases = {str(s.get("numero_cnj")) for s in list_accepted_sources(root)}
    candidates: list[dict[str, object]] = []
    for record in list_feedback(root):
        if not record.get("corpus_usable"):
            continue
        numero_cnj = str(record.get("numero_cnj") or "")
        candidates.append(
            {
                "numero_cnj": numero_cnj,
                "output_dir": record.get("output_dir"),
                "missing_source": record.get("missing_source"),
                "notes": record.get("notes"),
                "created_at": record.get("created_at"),
                "accepted": numero_cnj in accepted_cases,
            }
        )
    return candidates


def _require_provenance(payload: dict[str, object]) -> None:
    """Every corpus source must carry full provenance (URL/fonte/data/hash/tipo).

    URL and hash are enforced by _public_source_url / _resolve_content_hash; here we
    require the type, date, and a fonte (tribunal or publisher) so nothing enters the
    moat without an auditable origin.
    """
    for field in ("source_type", "source_date"):
        if not str(payload.get(field) or "").strip():
            msg = f"{field} é obrigatório para proveniência da fonte."
            raise ValueError(msg)
    if not (str(payload.get("tribunal") or "").strip() or str(payload.get("source_publisher") or "").strip()):
        msg = "fonte é obrigatória (tribunal ou source_publisher) para proveniência."
        raise ValueError(msg)
    kind = str(payload.get("provenance_kind") or "publica")
    if kind not in {"publica", "acervo_do_escritorio"}:
        msg = "provenance_kind deve ser 'publica' ou 'acervo_do_escritorio'."
        raise ValueError(msg)
    tipo_raw = str(payload.get("source_type") or "")
    if tipo_raw in {TipoFonte.DOUTRINA_PD.value, TipoFonte.DOUTRINA_PRIVADA.value}:
        rights = str(payload.get("rights_basis") or "")
        if rights not in RIGHTS_BASIS_VALUES:
            msg = (
                "rights_basis é obrigatório para doutrina "
                f"({', '.join(sorted(RIGHTS_BASIS_VALUES))}) — sem base de direitos não ingere."
            )
            raise ValueError(msg)
    override = str(payload.get("uso") or "")
    if override:
        resolve_uso(tipo_raw or None, override)  # ValueError se inválido


def append_accepted_source(root: Path, payload: dict[str, object]) -> dict[str, object]:
    """Append one lawyer-approved source with mandatory provenance."""
    ensure_private_dir(root)
    _require_provenance(payload)
    content_hash = _resolve_content_hash(payload)
    kind = str(payload.get("provenance_kind") or "publica")
    if kind == "acervo_do_escritorio":
        raw_url = str(payload.get("source_url") or "").strip()
        source_url = _public_source_url(raw_url) if raw_url else ""
    else:
        source_url = _public_source_url(payload.get("source_url"))
    if any(str(record.get("content_sha256") or "") == content_hash for record in list_accepted_sources(root)):
        msg = "fonte já aceita no corpus com o mesmo content_sha256."
        raise ValueError(msg)
    source_id = uuid.uuid4().hex
    text_path = _write_source_text(root, source_id, payload)
    record = {
        **{
            k: v
            for k, v in payload.items()
            if k
            not in {
                "id",
                "created_at",
                "content_sha256",
                "reingest_status",
                "reingested_at",
                "source_text",
                "source_text_path",
            }
        },
        "id": source_id,
        "created_at": datetime.now(UTC).isoformat(),
        "content_sha256": content_hash,
        "reingest_status": "pending",
        "source_text_path": text_path,
        "source_url": source_url,
        "provenance_kind": kind,
        "uso": resolve_uso(
            str(payload.get("source_type") or "") or None, str(payload.get("uso") or "") or None
        ).value,
        "tipo_peticao": str(payload.get("tipo_peticao") or ""),
        "rights_basis": str(payload.get("rights_basis") or ""),
    }
    with sources_path(root).open("a", encoding="utf-8") as fh:
        fh.write(json.dumps(record, ensure_ascii=False, sort_keys=True) + "\n")
    restrict_file(sources_path(root))
    return record


def list_accepted_sources(root: Path) -> list[dict[str, object]]:
    path = sources_path(root)
    if not path.exists():
        return []
    records: list[dict[str, object]] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        try:
            data = json.loads(line)
        except json.JSONDecodeError:
            continue
        if isinstance(data, dict):
            records.append(data)
    records.sort(key=lambda r: str(r.get("created_at", "")), reverse=True)
    return records


def coverage_report(root: Path) -> dict[str, object]:
    """Coverage and controlled reingestion report for accepted sources."""
    sources = list_accepted_sources(root)
    candidates = corpus_candidates(root)
    return {
        "accepted_count": len(sources),
        "pending_candidates": [c for c in candidates if not c.get("accepted")],
        "pending_reingest": [s for s in sources if s.get("reingest_status") == "pending"],
        "coverage": {
            "area": _counts(sources, "area"),
            "tema": _counts(sources, "tema"),
            "tribunal": _counts(sources, "tribunal"),
            "source_type": _counts(sources, "source_type"),
            "status": _counts(sources, "status"),
        },
    }


def mark_reingested(root: Path, source_id: str) -> dict[str, object] | None:
    """Mark an accepted source as reingested after the controlled corpus job runs."""
    records = list_accepted_sources(root)
    updated: dict[str, object] | None = None
    now = datetime.now(UTC).isoformat()
    for record in records:
        if record.get("id") == source_id:
            record["reingest_status"] = "done"
            record["reingested_at"] = now
            updated = record
            break
    if updated is None:
        return None
    _rewrite_sources(root, records)
    return updated


def reingest_pending_sources(
    root: Path, repertory_path: Path, tenant_id: str | None = None
) -> ReingestReport:
    """Ingest pending accepted sources into the local FTS repertory.

    ``tenant_id`` tags the uploaded chunks so a firm's private corpus (tier-2/3) is only
    retrievable by that firm — never commingled with another tenant's search.
    """
    from juris.repertory.chunking import chunk_fonte
    from juris.repertory.corpus.models import TIPO_HIERARQUIA, FonteJurisprudencia, TipoFonte
    from juris.repertory.vector_store import LocalFTSStore

    records = list_accepted_sources(root)
    pending = [r for r in records if r.get("reingest_status") == "pending"]
    if not pending:
        return ReingestReport(processed=0, chunks=0, errors=[])

    store = LocalFTSStore(repertory_path)
    processed = 0
    total_chunks = 0
    errors: list[dict[str, str]] = []
    for record in pending:
        source_id = str(record.get("id") or "")
        try:
            text = _read_source_text(root, record)
            tipo = TipoFonte(str(record.get("source_type") or "acordao_publicado"))
            fonte = FonteJurisprudencia(
                id=f"pilot-{source_id}",
                tribunal=str(record.get("tribunal") or ""),
                tipo=tipo,
                numero=str(record.get("title") or record.get("numero_cnj") or source_id),
                ementa=str(record.get("title") or ""),
                texto_integral=text,
                data_julgamento=_parse_date(str(record.get("source_date") or "")),
                temas=[str(record.get("tema") or "")],
                situacao=str(record.get("status") or "vigente"),
                hierarquia=TIPO_HIERARQUIA.get(tipo, 6),
                source_url=str(record.get("source_url") or ""),
            )
            chunks = chunk_fonte(fonte)
            uso_val = str(record.get("uso") or "") or resolve_uso(tipo).value
            resolved_chunks = []
            for chunk in chunks:
                # DocumentChunk is frozen/slots — dataclasses.replace to set uso;
                # metadata stays the same dict reference, so updating it in place is fine.
                chunk = replace(chunk, uso=uso_val)
                chunk.metadata.update(
                    {
                        "pilot_source_id": source_id,
                        "numero_cnj": record.get("numero_cnj"),
                        "source_url": record.get("source_url"),
                        "source_date": record.get("source_date"),
                        "content_sha256": record.get("content_sha256"),
                        "area": record.get("area"),
                        "tema": record.get("tema"),
                        "tipo_peticao": record.get("tipo_peticao"),
                    }
                )
                resolved_chunks.append(chunk)
            chunks = resolved_chunks
            stored = store.upsert(chunks, [[] for _ in chunks], tenant_id=tenant_id)
            mark_reingested(root, source_id)
            processed += 1
            total_chunks += stored
        except Exception as exc:  # noqa: BLE001 - report per-source failures
            from juris.core.observability import get_logger
            from juris.core.sanitize import safe_error_text

            get_logger("juris.web").warning(
                "corpus_reingest_source_error",
                source_id=source_id,
                error=safe_error_text(exc),
                exception_type=exc.__class__.__name__,
            )
            errors.append({"source_id": source_id, "error": _REINGEST_SOURCE_ERROR})
    return ReingestReport(processed=processed, chunks=total_chunks, errors=errors)


_UPLOAD_MAX_BYTES = 20 * 1024 * 1024
_UPLOAD_MAX_CHARS = 2_000_000


def extract_upload_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded office document (.pdf, .docx, .txt or .md).

    The office-archive upload is the ToS-approved inteiro-teor path
    (``data/tos_compliance_log.md``): documents the firm already owns.
    """
    if len(data) > _UPLOAD_MAX_BYTES:
        msg = "arquivo excede o limite de 20MB."
        raise ValueError(msg)
    name = filename.lower().strip()
    if name.endswith(".pdf"):
        import pymupdf  # heavy import kept off the module path

        try:
            doc = cast("Any", pymupdf.open(stream=data, filetype="pdf"))  # type: ignore[no-untyped-call]
            with doc:
                text = "\n".join(page.get_text() for page in doc)
        except (RuntimeError, ValueError) as exc:
            msg = "não foi possível ler o PDF — exporte novamente ou cole o texto."
            raise ValueError(msg) from exc
    elif name.endswith(".docx"):
        import io
        from zipfile import BadZipFile

        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = Document(io.BytesIO(data))
        except (PackageNotFoundError, KeyError, ValueError, BadZipFile) as exc:
            msg = "não foi possível ler o DOCX — exporte novamente ou cole o texto."
            raise ValueError(msg) from exc
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells if cell.text.strip())
        text = "\n".join(parts)
    elif name.endswith((".txt", ".md")):
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            # Exportações jurídicas brasileiras são frequentemente cp1252.
            text = data.decode("cp1252", errors="replace")
    else:
        msg = "formato não suportado — envie PDF, DOCX, TXT ou MD, ou cole o texto."
        raise ValueError(msg)
    text = text.strip()
    if not text:
        msg = "documento sem texto extraível — cole o texto manualmente."
        raise ValueError(msg)
    return text[:_UPLOAD_MAX_CHARS]


def upload_source_document(
    root: Path, repertory_path: Path, payload: dict[str, object], tenant_id: str | None
) -> dict[str, object]:
    """Register one office-archive document and ingest it immediately.

    Combines the two existing controlled steps (accepted source with mandatory
    provenance + pending reingestion) into the single flow the console exposes.
    """
    text = str(payload.get("source_text") or "").strip()
    encoded = str(payload.get("content_base64") or "").strip()
    if not text and encoded:
        try:
            data = base64.b64decode(encoded, validate=True)
        except binascii.Error as exc:
            msg = "content_base64 inválido."
            raise ValueError(msg) from exc
        text = extract_upload_text(str(payload.get("filename") or ""), data)
    if not text:
        msg = "envie o texto da decisão (source_text) ou um arquivo (filename + content_base64)."
        raise ValueError(msg)

    record_payload: dict[str, object] = {
        key: payload[key]
        for key in (
            "title",
            "source_type",
            "source_date",
            "source_url",
            "tribunal",
            "source_publisher",
            "tema",
            "area",
            "numero_cnj",
            "provenance_kind",
            "uso",
            "tipo_peticao",
            "rights_basis",
        )
        if str(payload.get(key) or "").strip()
    }
    record_payload["source_text"] = text
    source = append_accepted_source(root, record_payload)
    report = reingest_pending_sources(root, repertory_path, tenant_id=tenant_id)
    return {"source": source, "reingest": report.to_dict()}


def _resolve_content_hash(payload: dict[str, object]) -> str:
    text = str(payload.get("source_text") or "").strip()
    explicit = str(payload.get("content_sha256") or "").strip().lower()
    if explicit:
        if _SHA256_RE.fullmatch(explicit) is None:
            msg = "content_sha256 deve ser um SHA-256 hexadecimal de 64 caracteres."
            raise ValueError(msg)
        if text:
            computed = hashlib.sha256(text.encode("utf-8")).hexdigest()
            if explicit != computed:
                msg = "content_sha256 não confere com source_text."
                raise ValueError(msg)
        return explicit
    if not text:
        msg = "content_sha256 ou source_text é obrigatório para proveniência."
        raise ValueError(msg)
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _public_source_url(value: object) -> str:
    """Return a provenance URL safe for UI/export/corpus metadata."""
    raw = str(value or "").strip()
    parsed = urlparse(raw)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        msg = "source_url deve ser uma URL http(s)."
        raise ValueError(msg)
    if parsed.username or parsed.password:
        msg = "source_url não pode conter credenciais."
        raise ValueError(msg)
    return urlunparse((parsed.scheme, parsed.netloc, parsed.path, parsed.params, parsed.query, ""))


def _write_source_text(root: Path, source_id: str, payload: dict[str, object]) -> str | None:
    text = str(payload.get("source_text") or "").strip()
    if not text:
        return None
    text_dir = root / _TEXT_DIR
    ensure_private_dir(text_dir)
    path = text_dir / f"{source_id}.txt"
    path.write_text(text, encoding="utf-8")
    restrict_file(path)
    return str(path.relative_to(root))


def _read_source_text(root: Path, record: dict[str, object]) -> str:
    rel_path = record.get("source_text_path")
    if not rel_path:
        msg = "fonte sem texto local; reingestão automática ainda não disponível só por URL/hash."
        raise RuntimeError(msg)
    path = (root / str(rel_path)).resolve()
    if not path.is_relative_to(root.resolve()):
        msg = "source_text_path fora do diretório do tenant."
        raise RuntimeError(msg)
    return path.read_text(encoding="utf-8")


def _parse_date(value: str) -> date | None:
    try:
        return date.fromisoformat(value)
    except ValueError:
        return None


def _counts(records: list[dict[str, object]], key: str) -> dict[str, int]:
    counts: dict[str, int] = {}
    for record in records:
        value = str(record.get(key) or "").strip() or "não informado"
        counts[value] = counts.get(value, 0) + 1
    return dict(sorted(counts.items(), key=lambda item: (-item[1], item[0])))


def _rewrite_sources(root: Path, records: list[dict[str, object]]) -> None:
    ensure_private_dir(root)
    path = sources_path(root)
    with path.open("w", encoding="utf-8") as fh:
        for record in sorted(records, key=lambda r: str(r.get("created_at", ""))):
            fh.write(json.dumps(record, ensure_ascii=False, sort_keys=True) + "\n")
    restrict_file(path)

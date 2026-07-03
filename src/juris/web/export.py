"""Export a Markdown minuta to .docx so the lawyer can finish it in Word.

Deliberately small: legal drafts use headings, paragraphs, bold, and simple
bullet/numbered lists. We don't pull a full Markdown engine — the parser here
covers exactly those and stays auditable.
"""

from __future__ import annotations

import io
import re

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*|__(.+?)__")


def _add_runs_with_bold(paragraph: object, text: str) -> None:
    """Split ``text`` on **bold**/__bold__ spans, adding bold/plain runs."""
    para = paragraph  # typed as object to avoid importing docx types at module load
    pos = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > pos:
            para.add_run(text[pos : match.start()])  # type: ignore[attr-defined]
        run = para.add_run(match.group(1) or match.group(2) or "")  # type: ignore[attr-defined]
        run.bold = True
        pos = match.end()
    if pos < len(text):
        para.add_run(text[pos:])  # type: ignore[attr-defined]


def markdown_to_docx(markdown: str, *, title: str | None = None) -> bytes:
    """Render a Markdown minuta to a .docx byte stream (headings/paragraphs/bold/lists)."""
    from docx import Document  # heavy import kept off the module path

    document = Document()
    if title:
        document.core_properties.title = title

    lines = markdown.replace("\r\n", "\n").split("\n")
    paragraph_buffer: list[str] = []

    def _flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        paragraph_buffer.clear()
        if text:
            _add_runs_with_bold(document.add_paragraph(), text)

    for line in lines:
        heading = _HEADING_RE.match(line)
        bullet = _BULLET_RE.match(line)
        numbered = _NUMBERED_RE.match(line)
        if heading:
            _flush_paragraph()
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2).strip(), level=level)
        elif bullet:
            _flush_paragraph()
            _add_runs_with_bold(document.add_paragraph(style="List Bullet"), bullet.group(1).strip())
        elif numbered:
            _flush_paragraph()
            _add_runs_with_bold(document.add_paragraph(style="List Number"), numbered.group(1).strip())
        elif not line.strip():
            _flush_paragraph()
        else:
            paragraph_buffer.append(line)
    _flush_paragraph()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
